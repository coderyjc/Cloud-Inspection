from docx import Document

doc = Document('./JancoBlog个人博客源程序V2.0.docx')

target = Document()

for p in doc.paragraphs:
  if p.text.startswith('R:'):
    target.add_heading(p.text, level=3)
    print(p.text)
  elif p.text.startswith('R:'):
    target.add_heading(p.text, level=3)
    print(p.text)
  else:
    target.add_paragraph(p.text)

target.save("target.docx")